"""
Document Generator module for creating Word documents from completed canvases.
"""

import threading

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from io import BytesIO
from typing import List
from datetime import datetime


class DocumentGenerator:
    """Generates professional Word documents from Value Proposition Canvas data."""

    # Color scheme
    PRIMARY_COLOR = RGBColor(79, 70, 229)  # Indigo
    SUCCESS_COLOR = RGBColor(16, 185, 129)  # Emerald
    WARNING_COLOR = RGBColor(245, 158, 11)  # Amber
    TEXT_COLOR = RGBColor(31, 41, 55)  # Gray-800

    def __init__(self):
        self.document = None
        self._lock = threading.Lock()
    
    def _setup_styles(self):
        """Set up custom styles for the document."""
        styles = self.document.styles
        
        # Title style
        if 'Canvas Title' not in [s.name for s in styles]:
            title_style = styles.add_style('Canvas Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(28)
            title_font.bold = True
            title_font.color.rgb = self.PRIMARY_COLOR
        
        # Section heading style
        if 'Section Heading' not in [s.name for s in styles]:
            section_style = styles.add_style('Section Heading', WD_STYLE_TYPE.PARAGRAPH)
            section_font = section_style.font
            section_font.name = 'Calibri'
            section_font.size = Pt(16)
            section_font.bold = True
            section_font.color.rgb = self.PRIMARY_COLOR
        
        # Body text style
        if 'Canvas Body' not in [s.name for s in styles]:
            body_style = styles.add_style('Canvas Body', WD_STYLE_TYPE.PARAGRAPH)
            body_font = body_style.font
            body_font.name = 'Calibri'
            body_font.size = Pt(11)
            body_font.color.rgb = self.TEXT_COLOR
    
    def _add_title(self, title: str):
        """Add the document title."""
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(title)
        run.font.name = 'Calibri'
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = self.PRIMARY_COLOR
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.space_after = Pt(6)
    
    def _add_subtitle(self, text: str):
        """Add a subtitle."""
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(107, 114, 128)  # Gray-500
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.space_after = Pt(24)
    
    def _add_section_heading(self, title: str, emoji: str = ""):
        """Add a section heading."""
        self.document.add_paragraph()  # Spacer
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(f"{emoji} {title}" if emoji else title)
        run.font.name = 'Calibri'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = self.PRIMARY_COLOR
        paragraph.space_after = Pt(12)
    
    def _add_description_box(self, text: str):
        """Add a boxed description (simulated with indentation and styling)."""
        paragraph = self.document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.right_indent = Inches(0.25)
        run = paragraph.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.color.rgb = self.TEXT_COLOR
        run.font.italic = True
        paragraph.space_after = Pt(16)
    
    def _add_numbered_list(self, items: List[str], color: RGBColor = None):
        """Add a numbered list of items."""
        for i, item in enumerate(items, 1):
            paragraph = self.document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            
            # Number
            num_run = paragraph.add_run(f"{i}. ")
            num_run.font.name = 'Calibri'
            num_run.font.size = Pt(11)
            num_run.font.bold = True
            num_run.font.color.rgb = color or self.PRIMARY_COLOR
            
            # Text
            text_run = paragraph.add_run(item)
            text_run.font.name = 'Calibri'
            text_run.font.size = Pt(11)
            text_run.font.color.rgb = self.TEXT_COLOR
            
            paragraph.space_after = Pt(8)
    
    def _add_footer(self):
        """Add a footer with generation date."""
        self.document.add_paragraph()
        self.document.add_paragraph()
        
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run("─" * 60)
        run.font.color.rgb = RGBColor(209, 213, 219)  # Gray-300
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        paragraph = self.document.add_paragraph()
        date_str = datetime.now().strftime("%B %d, %Y at %H:%M")
        run = paragraph.add_run(f"Generated on {date_str}")
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(156, 163, 175)  # Gray-400
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run("Work Process Reflection Canvas")
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(156, 163, 175)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def generate(self, job_description: str, pain_points: List[str],
                 gain_points: List[str], title: str = "Work Process Canvas") -> BytesIO:
        """
        Generate a Word document from the canvas data.
        Thread-safe: uses a lock to prevent concurrent access to self.document.
        """
        with self._lock:
            return self._generate_impl(job_description, pain_points, gain_points, title)

    def _generate_impl(self, job_description: str, pain_points: List[str],
                       gain_points: List[str], title: str) -> BytesIO:
        self.document = Document()
        
        # Set up styles
        self._setup_styles()
        
        # Add title
        self._add_title(title)
        self._add_subtitle("Work Process Analysis")
        
        # Job Description section
        self._add_section_heading("Job Description", "🎯")
        self._add_description_box(job_description)
        
        # Pain Points section
        self._add_section_heading(f"Pain Points ({len(pain_points)} identified)", "😓")
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run("Obstacles, frustrations, and risks in your work:")
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(107, 114, 128)
        paragraph.space_after = Pt(8)
        self._add_numbered_list(pain_points, self.WARNING_COLOR)
        
        # Gain Points section
        self._add_section_heading(f"Gain Points ({len(gain_points)} identified)", "🌟")
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run("Outcomes and benefits you desire:")
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(107, 114, 128)
        paragraph.space_after = Pt(8)
        self._add_numbered_list(gain_points, self.SUCCESS_COLOR)
        
        # Summary section
        self._add_section_heading("Summary", "📋")
        summary_text = f"""This Work Process Canvas captures your work profile with:
• 1 clearly defined job description
• {len(pain_points)} distinct pain points
• {len(gain_points)} unique gain points

Use this analysis to identify improvements and optimize your work processes."""
        
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(summary_text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = self.TEXT_COLOR
        
        # Footer
        self._add_footer()
        
        # Save to BytesIO
        buffer = BytesIO()
        self.document.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    def generate_to_file(self, job_description: str, pain_points: List[str],
                         gain_points: List[str], filepath: str,
                         title: str = "Work Process Canvas"):
        """
        Generate and save a Word document to a file.
        
        Args:
            job_description: The job description
            pain_points: List of pain points
            gain_points: List of gain points
            filepath: Path to save the document
            title: Document title
        """
        buffer = self.generate(job_description, pain_points, gain_points, title)
        
        with open(filepath, 'wb') as f:
            f.write(buffer.getvalue())
